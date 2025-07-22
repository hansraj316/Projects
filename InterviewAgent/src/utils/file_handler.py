"""
File handling utilities for InterviewAgent
"""

import os
import io
import logging
from typing import Dict, Any, Optional, Union, BinaryIO
from pathlib import Path
import tempfile

# PDF and document processing imports
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

PDF_AVAILABLE = PYPDF_AVAILABLE or PYPDF2_AVAILABLE

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class FileHandler:
    """Handle file upload, validation, and storage operations"""
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self, upload_dir: Optional[str] = None):
        """Initialize file handler with upload directory"""
        if upload_dir:
            self.upload_dir = Path(upload_dir)
        else:
            # Use temp directory if none specified
            self.upload_dir = Path(tempfile.gettempdir()) / "interviewagent_uploads"
        
        # Create upload directory if it doesn't exist
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"File handler initialized with upload dir: {self.upload_dir}")
    
    def validate_file(self, file_obj, filename: str) -> Dict[str, Any]:
        """
        Validate uploaded file
        
        Args:
            file_obj: File object or bytes
            filename: Original filename
            
        Returns:
            Dict with validation result
        """
        try:
            # Check file extension
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.ALLOWED_EXTENSIONS:
                return {
                    "valid": False,
                    "error": f"File type {file_ext} not supported. Allowed types: {', '.join(self.ALLOWED_EXTENSIONS)}"
                }
            
            # Check file size
            if hasattr(file_obj, 'size'):
                file_size = file_obj.size
            elif hasattr(file_obj, 'seek') and hasattr(file_obj, 'tell'):
                # Get size by seeking
                current_pos = file_obj.tell()
                file_obj.seek(0, 2)  # Seek to end
                file_size = file_obj.tell()
                file_obj.seek(current_pos)  # Reset position
            else:
                # For bytes objects
                file_size = len(file_obj) if isinstance(file_obj, (bytes, bytearray)) else 0
            
            if file_size > self.MAX_FILE_SIZE:
                return {
                    "valid": False,
                    "error": f"File size ({file_size / (1024*1024):.2f}MB) exceeds maximum allowed size ({self.MAX_FILE_SIZE / (1024*1024)}MB)"
                }
            
            # Basic content validation
            try:
                if hasattr(file_obj, 'read'):
                    content_preview = file_obj.read(1024)  # Read first 1KB
                    if hasattr(file_obj, 'seek'):
                        file_obj.seek(0)  # Reset for later reading
                else:
                    content_preview = file_obj[:1024] if isinstance(file_obj, (bytes, bytearray)) else b""
                
                # Check if file is empty
                if not content_preview:
                    return {
                        "valid": False,
                        "error": "File appears to be empty"
                    }
                
            except Exception as e:
                logger.warning(f"Could not validate file content: {e}")
            
            return {
                "valid": True,
                "file_size": file_size,
                "file_extension": file_ext,
                "filename": filename
            }
            
        except Exception as e:
            logger.error(f"File validation error: {e}")
            return {
                "valid": False,
                "error": f"File validation failed: {str(e)}"
            }
    
    def save_file(self, file_obj, filename: str, user_id: str = "default") -> Dict[str, Any]:
        """
        Save uploaded file to storage
        
        Args:
            file_obj: File object or bytes
            filename: Original filename
            user_id: User identifier
            
        Returns:
            Dict with save result and file info
        """
        try:
            # Validate file first
            validation = self.validate_file(file_obj, filename)
            if not validation["valid"]:
                return validation
            
            # Create user directory
            user_dir = self.upload_dir / user_id
            user_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate safe filename with timestamp
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_ext = Path(filename).suffix.lower()
            safe_filename = f"{timestamp}_{self._sanitize_filename(Path(filename).stem)}{file_ext}"
            file_path = user_dir / safe_filename
            
            # Save file
            if hasattr(file_obj, 'read'):
                # File-like object
                content = file_obj.read()
            else:
                # Assume bytes
                content = file_obj
            
            with open(file_path, 'wb') as f:
                f.write(content)
            
            logger.info(f"File saved: {file_path}")
            
            return {
                "success": True,
                "file_path": str(file_path),
                "filename": safe_filename,
                "original_filename": filename,
                "file_size": validation["file_size"],
                "file_extension": validation["file_extension"],
                "user_id": user_id,
                "saved_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"File save error: {e}")
            return {
                "success": False,
                "error": f"Failed to save file: {str(e)}"
            }
    
    def delete_file(self, file_path: str) -> bool:
        """Delete a file safely"""
        try:
            path = Path(file_path)
            if path.exists() and path.is_file():
                # Ensure the file is within our upload directory
                if str(path.resolve()).startswith(str(self.upload_dir.resolve())):
                    path.unlink()
                    logger.info(f"File deleted: {file_path}")
                    return True
                else:
                    logger.warning(f"Attempted to delete file outside upload directory: {file_path}")
                    return False
            return False
        except Exception as e:
            logger.error(f"File deletion error: {e}")
            return False
    
    def list_user_files(self, user_id: str) -> list:
        """List all files for a user"""
        try:
            user_dir = self.upload_dir / user_id
            if not user_dir.exists():
                return []
            
            files = []
            for file_path in user_dir.glob("*"):
                if file_path.is_file():
                    try:
                        stat = file_path.stat()
                        files.append({
                            "filename": file_path.name,
                            "file_path": str(file_path),
                            "size": stat.st_size,
                            "created_at": stat.st_ctime,
                            "modified_at": stat.st_mtime
                        })
                    except Exception as e:
                        logger.warning(f"Could not get file stats for {file_path}: {e}")
            
            return sorted(files, key=lambda x: x["modified_at"], reverse=True)
            
        except Exception as e:
            logger.error(f"Error listing user files: {e}")
            return []
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to prevent path traversal and invalid characters"""
        # Remove or replace dangerous characters
        import re
        # Keep only alphanumeric, spaces, hyphens, underscores, and periods
        safe_name = re.sub(r'[^a-zA-Z0-9\s\-_.]', '_', filename)
        # Replace multiple spaces/underscores with single underscore
        safe_name = re.sub(r'[\s_]+', '_', safe_name)
        # Limit length
        return safe_name[:100]
    
    def get_file_info(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Get information about a file"""
        try:
            path = Path(file_path)
            if not path.exists():
                return None
            
            stat = path.stat()
            return {
                "filename": path.name,
                "file_path": str(path),
                "size": stat.st_size,
                "extension": path.suffix.lower(),
                "created_at": stat.st_ctime,
                "modified_at": stat.st_mtime,
                "exists": True
            }
            
        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return None


def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Extract text content from various file types
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dict with extracted text and metadata
    """
    try:
        path = Path(file_path)
        file_ext = path.suffix.lower()
        
        if file_ext == '.txt':
            return _extract_text_from_txt(file_path)
        elif file_ext == '.pdf':
            return _extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return _extract_text_from_docx(file_path)
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_ext}"
            }
            
    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        return {
            "success": False,
            "error": f"Failed to extract text: {str(e)}"
        }


def _extract_text_from_txt(file_path: str) -> Dict[str, Any]:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            "success": True,
            "text": content,
            "word_count": len(content.split()),
            "char_count": len(content),
            "file_type": "txt"
        }
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()
            return {
                "success": True,
                "text": content,
                "word_count": len(content.split()),
                "char_count": len(content),
                "file_type": "txt",
                "encoding": "latin-1"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to read TXT file: {str(e)}"
            }
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to extract text from TXT: {str(e)}"
        }


def _extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """Extract text from PDF file using pypdf (preferred) or PyPDF2 (fallback)"""
    if not PDF_AVAILABLE:
        return {
            "success": False,
            "error": "PDF processing not available. Install pypdf: pip install pypdf"
        }
    
    # Try pypdf first (more modern and reliable)
    if PYPDF_AVAILABLE:
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content.append(page_text.strip())
                    except Exception as e:
                        logger.warning(f"Could not extract text from page {page_num} with pypdf: {e}")
                
                full_text = "\n\n".join(text_content)
                
                # Validate extraction
                if not full_text.strip():
                    raise Exception("No text content extracted from PDF")
                
                return {
                    "success": True,
                    "text": full_text,
                    "word_count": len(full_text.split()),
                    "char_count": len(full_text),
                    "page_count": len(pdf_reader.pages),
                    "file_type": "pdf",
                    "extraction_library": "pypdf"
                }
                
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}. Trying PyPDF2 fallback...")
    
    # Fallback to PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content.append(page_text.strip())
                    except Exception as e:
                        logger.warning(f"Could not extract text from page {page_num} with PyPDF2: {e}")
                
                full_text = "\n\n".join(text_content)
                
                # Validate extraction
                if not full_text.strip():
                    raise Exception("No text content extracted from PDF")
                
                return {
                    "success": True,
                    "text": full_text,
                    "word_count": len(full_text.split()),
                    "char_count": len(full_text),
                    "page_count": len(pdf_reader.pages),
                    "file_type": "pdf",
                    "extraction_library": "PyPDF2"
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to extract text from PDF with both pypdf and PyPDF2: {str(e)}"
            }
    
    return {
        "success": False,
        "error": "No PDF processing library available"
    }


def _extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        return {
            "success": False,
            "error": "DOCX processing not available. Install python-docx: pip install python-docx"
        }
    
    try:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = []
        all_text.extend(paragraphs)
        if table_text:
            all_text.append("\n--- Tables ---")
            all_text.extend(table_text)
        
        full_text = "\n\n".join(all_text)
        
        return {
            "success": True,
            "text": full_text,
            "word_count": len(full_text.split()),
            "char_count": len(full_text),
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "file_type": "docx"
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to extract text from DOCX: {str(e)}"
        }