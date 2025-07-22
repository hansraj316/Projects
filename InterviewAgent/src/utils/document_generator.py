"""
Document Generation Utilities - PDF and DOCX generation for resumes and cover letters
Enhanced for professional document creation with proper formatting
"""

import os
import tempfile
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
from pathlib import Path
import logging

# PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# DOCX generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class DocumentGenerator:
    """
    Generates professional PDF and DOCX documents for resumes and cover letters
    """
    
    def __init__(self, output_dir: str = None):
        """
        Initialize document generator
        
        Args:
            output_dir: Directory to save generated documents
        """
        self.output_dir = output_dir or tempfile.gettempdir()
        self.logger = logging.getLogger(__name__)
        
        # Ensure output directory exists
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
        
        # Check dependencies
        if not REPORTLAB_AVAILABLE:
            self.logger.warning("ReportLab not available. PDF generation will be limited.")
        if not PYTHON_DOCX_AVAILABLE:
            self.logger.warning("python-docx not available. DOCX generation will be limited.")
    
    def generate_resume_pdf(self, resume_data: Dict[str, Any], job_details: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Generate a professional PDF resume
        
        Args:
            resume_data: Resume information
            job_details: Optional job details for customization
            
        Returns:
            Generation result with file path
        """
        if not REPORTLAB_AVAILABLE:
            return self._generate_fallback_resume_text(resume_data, "pdf")
        
        try:
            # Generate filename
            candidate_name = resume_data.get('name', 'Unknown').replace(' ', '_')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"resume_{candidate_name}_{timestamp}.pdf"
            filepath = os.path.join(self.output_dir, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=12,
                alignment=TA_CENTER,
                textColor=colors.darkblue
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=6,
                spaceBefore=12,
                textColor=colors.darkblue,
                borderWidth=1,
                borderColor=colors.darkblue,
                borderPadding=3
            )
            
            # Build document content
            story = []
            
            # Header - Name and Contact
            story.append(Paragraph(resume_data.get('name', 'Unknown'), title_style))
            
            contact_info = []
            if resume_data.get('email'):
                contact_info.append(resume_data['email'])
            if resume_data.get('phone'):
                contact_info.append(resume_data['phone'])
            if resume_data.get('location'):
                contact_info.append(resume_data['location'])
            
            if contact_info:
                story.append(Paragraph(' | '.join(contact_info), styles['Normal']))
            
            story.append(Spacer(1, 12))
            
            # Professional Summary
            if resume_data.get('professional_summary'):
                story.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
                story.append(Paragraph(resume_data['professional_summary'], styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Skills
            if resume_data.get('skills'):
                story.append(Paragraph('TECHNICAL SKILLS', heading_style))
                skills_text = ', '.join(resume_data['skills']) if isinstance(resume_data['skills'], list) else str(resume_data['skills'])
                story.append(Paragraph(skills_text, styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Experience
            if resume_data.get('experience'):
                story.append(Paragraph('PROFESSIONAL EXPERIENCE', heading_style))
                
                for exp in resume_data['experience']:
                    if isinstance(exp, dict):
                        # Position and Company
                        position_company = f"<b>{exp.get('position', 'Unknown Position')}</b> - {exp.get('company', 'Unknown Company')}"
                        story.append(Paragraph(position_company, styles['Normal']))
                        
                        # Duration
                        if exp.get('duration'):
                            story.append(Paragraph(f"<i>{exp['duration']}</i>", styles['Normal']))
                        
                        # Achievements
                        if exp.get('achievements'):
                            for achievement in exp['achievements']:
                                story.append(Paragraph(f"• {achievement}", styles['Normal']))
                        
                        story.append(Spacer(1, 8))
            
            # Education
            if resume_data.get('education'):
                story.append(Paragraph('EDUCATION', heading_style))
                
                for edu in resume_data['education']:
                    if isinstance(edu, dict):
                        edu_text = f"<b>{edu.get('degree', 'Unknown Degree')}</b> - {edu.get('school', 'Unknown School')}"
                        if edu.get('year'):
                            edu_text += f" ({edu['year']})"
                        story.append(Paragraph(edu_text, styles['Normal']))
                        story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            self.logger.info(f"Generated PDF resume: {filepath}")
            
            return {
                "success": True,
                "file_path": filepath,
                "filename": filename,
                "file_type": "pdf",
                "file_size": os.path.getsize(filepath),
                "generation_method": "reportlab"
            }
            
        except Exception as e:
            self.logger.error(f"PDF generation failed: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": "Failed to generate PDF resume"
            }
    
    def generate_resume_docx(self, resume_data: Dict[str, Any], job_details: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Generate a professional DOCX resume
        
        Args:
            resume_data: Resume information
            job_details: Optional job details for customization
            
        Returns:
            Generation result with file path
        """
        if not PYTHON_DOCX_AVAILABLE:
            return self._generate_fallback_resume_text(resume_data, "docx")
        
        try:
            # Generate filename
            candidate_name = resume_data.get('name', 'Unknown').replace(' ', '_')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"resume_{candidate_name}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Create DOCX document
            doc = Document()
            
            # Set up styles
            styles = doc.styles
            
            # Header - Name
            header = doc.add_heading(resume_data.get('name', 'Unknown'), level=1)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact Information
            contact_info = []
            if resume_data.get('email'):
                contact_info.append(resume_data['email'])
            if resume_data.get('phone'):
                contact_info.append(resume_data['phone'])
            if resume_data.get('location'):
                contact_info.append(resume_data['location'])
            
            if contact_info:
                contact_para = doc.add_paragraph(' | '.join(contact_info))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Add space
            
            # Professional Summary
            if resume_data.get('professional_summary'):
                doc.add_heading('PROFESSIONAL SUMMARY', level=2)
                doc.add_paragraph(resume_data['professional_summary'])
            
            # Skills
            if resume_data.get('skills'):
                doc.add_heading('TECHNICAL SKILLS', level=2)
                skills_text = ', '.join(resume_data['skills']) if isinstance(resume_data['skills'], list) else str(resume_data['skills'])
                doc.add_paragraph(skills_text)
            
            # Experience
            if resume_data.get('experience'):
                doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
                
                for exp in resume_data['experience']:
                    if isinstance(exp, dict):
                        # Position and Company
                        position_para = doc.add_paragraph()
                        position_run = position_para.add_run(f"{exp.get('position', 'Unknown Position')} - {exp.get('company', 'Unknown Company')}")
                        position_run.bold = True
                        
                        # Duration
                        if exp.get('duration'):
                            duration_para = doc.add_paragraph(exp['duration'])
                            duration_para.runs[0].italic = True
                        
                        # Achievements
                        if exp.get('achievements'):
                            for achievement in exp['achievements']:
                                doc.add_paragraph(achievement, style='List Bullet')
                        
                        doc.add_paragraph()  # Add space between jobs
            
            # Education
            if resume_data.get('education'):
                doc.add_heading('EDUCATION', level=2)
                
                for edu in resume_data['education']:
                    if isinstance(edu, dict):
                        edu_para = doc.add_paragraph()
                        edu_run = edu_para.add_run(f"{edu.get('degree', 'Unknown Degree')} - {edu.get('school', 'Unknown School')}")
                        edu_run.bold = True
                        
                        if edu.get('year'):
                            year_para = doc.add_paragraph(f"Graduated: {edu['year']}")
                            year_para.runs[0].italic = True
            
            # Save document
            doc.save(filepath)
            
            self.logger.info(f"Generated DOCX resume: {filepath}")
            
            return {
                "success": True,
                "file_path": filepath,
                "filename": filename,
                "file_type": "docx",
                "file_size": os.path.getsize(filepath),
                "generation_method": "python-docx"
            }
            
        except Exception as e:
            self.logger.error(f"DOCX generation failed: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": "Failed to generate DOCX resume"
            }
    
    def generate_cover_letter_pdf(self, cover_letter_data: Dict[str, Any], job_details: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Generate a professional PDF cover letter
        
        Args:
            cover_letter_data: Cover letter content
            job_details: Job details for customization
            
        Returns:
            Generation result with file path
        """
        if not REPORTLAB_AVAILABLE:
            return self._generate_fallback_cover_letter_text(cover_letter_data, "pdf")
        
        try:
            # Generate filename
            company_name = job_details.get('company_name', 'Company') if job_details else 'Company'
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"cover_letter_{company_name.replace(' ', '_')}_{timestamp}.pdf"
            filepath = os.path.join(self.output_dir, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=72)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Custom styles for business letter
            header_style = ParagraphStyle(
                'HeaderStyle',
                parent=styles['Normal'],
                fontSize=12,
                alignment=TA_LEFT,
                spaceAfter=6
            )
            
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['Normal'],
                fontSize=12,
                alignment=TA_JUSTIFY,
                spaceAfter=12,
                leading=14
            )
            
            # Build document content
            story = []
            
            # Header with candidate info
            if isinstance(cover_letter_data, dict) and cover_letter_data.get('header'):
                for header_line in cover_letter_data['header']:
                    story.append(Paragraph(header_line, header_style))
            
            story.append(Spacer(1, 24))  # Date space
            
            # Date
            today = datetime.now().strftime('%B %d, %Y')
            story.append(Paragraph(today, styles['Normal']))
            story.append(Spacer(1, 24))
            
            # Recipient info
            if job_details:
                if job_details.get('company_name'):
                    story.append(Paragraph(job_details['company_name'], styles['Normal']))
                if job_details.get('hiring_manager'):
                    story.append(Paragraph(job_details['hiring_manager'], styles['Normal']))
            
            story.append(Spacer(1, 24))
            
            # Salutation
            if isinstance(cover_letter_data, dict):
                salutation = cover_letter_data.get('salutation', 'Dear Hiring Manager,')
            else:
                salutation = 'Dear Hiring Manager,'
            
            story.append(Paragraph(salutation, styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Body paragraphs
            if isinstance(cover_letter_data, dict) and cover_letter_data.get('body_paragraphs'):
                for paragraph in cover_letter_data['body_paragraphs']:
                    story.append(Paragraph(paragraph, body_style))
            elif isinstance(cover_letter_data, dict) and cover_letter_data.get('full_text'):
                # Split full text into paragraphs
                paragraphs = cover_letter_data['full_text'].split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        story.append(Paragraph(paragraph.strip(), body_style))
            
            # Closing
            if isinstance(cover_letter_data, dict):
                closing = cover_letter_data.get('closing', 'Sincerely,')
                signature = cover_letter_data.get('signature', 'Your Name')
            else:
                closing = 'Sincerely,'
                signature = 'Your Name'
            
            story.append(Spacer(1, 12))
            story.append(Paragraph(closing, styles['Normal']))
            story.append(Spacer(1, 24))
            story.append(Paragraph(signature, styles['Normal']))
            
            # Build PDF
            doc.build(story)
            
            self.logger.info(f"Generated PDF cover letter: {filepath}")
            
            return {
                "success": True,
                "file_path": filepath,
                "filename": filename,
                "file_type": "pdf",
                "file_size": os.path.getsize(filepath),
                "generation_method": "reportlab"
            }
            
        except Exception as e:
            self.logger.error(f"Cover letter PDF generation failed: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": "Failed to generate PDF cover letter"
            }
    
    def generate_cover_letter_docx(self, cover_letter_data: Dict[str, Any], job_details: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Generate a professional DOCX cover letter
        
        Args:
            cover_letter_data: Cover letter content
            job_details: Job details for customization
            
        Returns:
            Generation result with file path
        """
        if not PYTHON_DOCX_AVAILABLE:
            return self._generate_fallback_cover_letter_text(cover_letter_data, "docx")
        
        try:
            # Generate filename
            company_name = job_details.get('company_name', 'Company') if job_details else 'Company'
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"cover_letter_{company_name.replace(' ', '_')}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Create DOCX document
            doc = Document()
            
            # Header with candidate info
            if isinstance(cover_letter_data, dict) and cover_letter_data.get('header'):
                for header_line in cover_letter_data['header']:
                    doc.add_paragraph(header_line)
            
            doc.add_paragraph()  # Space
            
            # Date
            today = datetime.now().strftime('%B %d, %Y')
            doc.add_paragraph(today)
            doc.add_paragraph()  # Space
            
            # Recipient info
            if job_details:
                if job_details.get('company_name'):
                    doc.add_paragraph(job_details['company_name'])
                if job_details.get('hiring_manager'):
                    doc.add_paragraph(job_details['hiring_manager'])
            
            doc.add_paragraph()  # Space
            
            # Salutation
            if isinstance(cover_letter_data, dict):
                salutation = cover_letter_data.get('salutation', 'Dear Hiring Manager,')
            else:
                salutation = 'Dear Hiring Manager,'
            
            doc.add_paragraph(salutation)
            doc.add_paragraph()  # Space
            
            # Body paragraphs
            if isinstance(cover_letter_data, dict) and cover_letter_data.get('body_paragraphs'):
                for paragraph in cover_letter_data['body_paragraphs']:
                    doc.add_paragraph(paragraph)
                    doc.add_paragraph()  # Space between paragraphs
            elif isinstance(cover_letter_data, dict) and cover_letter_data.get('full_text'):
                # Split full text into paragraphs
                paragraphs = cover_letter_data['full_text'].split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
                        doc.add_paragraph()  # Space between paragraphs
            
            # Closing
            if isinstance(cover_letter_data, dict):
                closing = cover_letter_data.get('closing', 'Sincerely,')
                signature = cover_letter_data.get('signature', 'Your Name')
            else:
                closing = 'Sincerely,'
                signature = 'Your Name'
            
            doc.add_paragraph(closing)
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph(signature)
            
            # Save document
            doc.save(filepath)
            
            self.logger.info(f"Generated DOCX cover letter: {filepath}")
            
            return {
                "success": True,
                "file_path": filepath,
                "filename": filename,
                "file_type": "docx",
                "file_size": os.path.getsize(filepath),
                "generation_method": "python-docx"
            }
            
        except Exception as e:
            self.logger.error(f"Cover letter DOCX generation failed: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": "Failed to generate DOCX cover letter"
            }
    
    def _generate_fallback_resume_text(self, resume_data: Dict[str, Any], format_type: str) -> Dict[str, Any]:
        """Generate fallback text resume when PDF/DOCX libraries are not available"""
        try:
            candidate_name = resume_data.get('name', 'Unknown').replace(' ', '_')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"resume_{candidate_name}_{timestamp}.txt"
            filepath = os.path.join(self.output_dir, filename)
            
            # Generate text content
            content = []
            content.append(f"RESUME - {resume_data.get('name', 'Unknown').upper()}")
            content.append("=" * 50)
            content.append("")
            
            # Contact info
            contact_info = []
            if resume_data.get('email'):
                contact_info.append(f"Email: {resume_data['email']}")
            if resume_data.get('phone'):
                contact_info.append(f"Phone: {resume_data['phone']}")
            if resume_data.get('location'):
                contact_info.append(f"Location: {resume_data['location']}")
            
            content.extend(contact_info)
            content.append("")
            
            # Professional Summary
            if resume_data.get('professional_summary'):
                content.append("PROFESSIONAL SUMMARY")
                content.append("-" * 20)
                content.append(resume_data['professional_summary'])
                content.append("")
            
            # Skills
            if resume_data.get('skills'):
                content.append("TECHNICAL SKILLS")
                content.append("-" * 15)
                skills_text = ', '.join(resume_data['skills']) if isinstance(resume_data['skills'], list) else str(resume_data['skills'])
                content.append(skills_text)
                content.append("")
            
            # Experience
            if resume_data.get('experience'):
                content.append("PROFESSIONAL EXPERIENCE")
                content.append("-" * 23)
                
                for exp in resume_data['experience']:
                    if isinstance(exp, dict):
                        content.append(f"{exp.get('position', 'Unknown Position')} - {exp.get('company', 'Unknown Company')}")
                        if exp.get('duration'):
                            content.append(f"Duration: {exp['duration']}")
                        
                        if exp.get('achievements'):
                            for achievement in exp['achievements']:
                                content.append(f"• {achievement}")
                        
                        content.append("")
            
            # Education
            if resume_data.get('education'):
                content.append("EDUCATION")
                content.append("-" * 9)
                
                for edu in resume_data['education']:
                    if isinstance(edu, dict):
                        edu_line = f"{edu.get('degree', 'Unknown Degree')} - {edu.get('school', 'Unknown School')}"
                        if edu.get('year'):
                            edu_line += f" ({edu['year']})"
                        content.append(edu_line)
                        content.append("")
            
            # Write to file
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(content))
            
            return {
                "success": True,
                "file_path": filepath,
                "filename": filename,
                "file_type": "txt",
                "file_size": os.path.getsize(filepath),
                "generation_method": "fallback_text",
                "note": f"Generated as text file because {format_type.upper()} libraries are not available"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to generate fallback {format_type} resume"
            }
    
    def _generate_fallback_cover_letter_text(self, cover_letter_data: Dict[str, Any], format_type: str) -> Dict[str, Any]:
        """Generate fallback text cover letter when PDF/DOCX libraries are not available"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"cover_letter_{timestamp}.txt"
            filepath = os.path.join(self.output_dir, filename)
            
            # Generate text content
            content = []
            content.append("COVER LETTER")
            content.append("=" * 12)
            content.append("")
            
            # Header
            if isinstance(cover_letter_data, dict) and cover_letter_data.get('header'):
                content.extend(cover_letter_data['header'])
                content.append("")
            
            # Date
            content.append(datetime.now().strftime('%B %d, %Y'))
            content.append("")
            
            # Body content
            if isinstance(cover_letter_data, dict):
                if cover_letter_data.get('salutation'):
                    content.append(cover_letter_data['salutation'])
                    content.append("")
                
                if cover_letter_data.get('body_paragraphs'):
                    for paragraph in cover_letter_data['body_paragraphs']:
                        content.append(paragraph)
                        content.append("")
                elif cover_letter_data.get('full_text'):
                    content.append(cover_letter_data['full_text'])
                    content.append("")
                
                if cover_letter_data.get('closing'):
                    content.append(cover_letter_data['closing'])
                    content.append("")
                    content.append("")
                
                if cover_letter_data.get('signature'):
                    content.append(cover_letter_data['signature'])
            
            # Write to file
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(content))
            
            return {
                "success": True,
                "file_path": filepath,
                "filename": filename,
                "file_type": "txt",
                "file_size": os.path.getsize(filepath),
                "generation_method": "fallback_text",
                "note": f"Generated as text file because {format_type.upper()} libraries are not available"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to generate fallback {format_type} cover letter"
            }
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats"""
        formats = ["txt"]  # Always supported
        
        if REPORTLAB_AVAILABLE:
            formats.append("pdf")
        
        if PYTHON_DOCX_AVAILABLE:
            formats.append("docx")
        
        return formats
    
    def check_dependencies(self) -> Dict[str, Any]:
        """Check which document generation dependencies are available"""
        return {
            "reportlab_available": REPORTLAB_AVAILABLE,
            "python_docx_available": PYTHON_DOCX_AVAILABLE,
            "supported_formats": self.get_supported_formats(),
            "recommendations": [
                "pip install reportlab" if not REPORTLAB_AVAILABLE else "ReportLab: ✅ Available",
                "pip install python-docx" if not PYTHON_DOCX_AVAILABLE else "python-docx: ✅ Available"
            ]
        }


def generate_document(document_type: str, data: Dict[str, Any], format_type: str = "pdf", 
                     output_dir: str = None, job_details: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Convenience function to generate documents
    
    Args:
        document_type: 'resume' or 'cover_letter'
        data: Document data
        format_type: 'pdf', 'docx', or 'txt'
        output_dir: Output directory
        job_details: Optional job details for customization
        
    Returns:
        Generation result
    """
    generator = DocumentGenerator(output_dir)
    
    if document_type == "resume":
        if format_type == "pdf":
            return generator.generate_resume_pdf(data, job_details)
        elif format_type == "docx":
            return generator.generate_resume_docx(data, job_details)
        else:
            return generator._generate_fallback_resume_text(data, format_type)
    
    elif document_type == "cover_letter":
        if format_type == "pdf":
            return generator.generate_cover_letter_pdf(data, job_details)
        elif format_type == "docx":
            return generator.generate_cover_letter_docx(data, job_details)
        else:
            return generator._generate_fallback_cover_letter_text(data, format_type)
    
    else:
        return {
            "success": False,
            "error": f"Unknown document type: {document_type}",
            "message": "Supported types: 'resume', 'cover_letter'"
        }